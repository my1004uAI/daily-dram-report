import os
import re
import urllib.parse
import datetime
from collections import defaultdict

import requests
import feedparser
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request as GoogleRequest
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# =========================
# 환경변수
# =========================
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET", "")
GOOGLE_REFRESH_TOKEN = os.getenv("GOOGLE_REFRESH_TOKEN", "")
GOOGLE_DRIVE_FOLDER_ID = os.getenv("GOOGLE_DRIVE_FOLDER_ID", "")

print("GOOGLE_CLIENT_ID exists:", bool(GOOGLE_CLIENT_ID))
print("GOOGLE_CLIENT_SECRET exists:", bool(GOOGLE_CLIENT_SECRET))
print("GOOGLE_REFRESH_TOKEN exists:", bool(GOOGLE_REFRESH_TOKEN))
print("GOOGLE_DRIVE_FOLDER_ID exists:", bool(GOOGLE_DRIVE_FOLDER_ID))

# =========================
# 기본 설정
# =========================
SEARCH_TERMS = ["DRAM", "HBM", "DDR5", "DDR6", "LPDDR5", "LPDDR6"]

COMPANY_KEYWORDS = {
    "Samsung": ["samsung", "samsung electronics"],
    "SK hynix": ["sk hynix", "hynix"],
    "Micron": ["micron", "micron technology"],
    "CXMT": ["cxmt", "changxin", "changxin memory", "changxin memory technologies"],
}

ARXIV_MAX_RESULTS = 30

KST = datetime.timezone(datetime.timedelta(hours=9))
now_kst = datetime.datetime.now(KST)
now_utc = datetime.datetime.utcnow()
cutoff_utc = now_utc - datetime.timedelta(days=1)

today_ymd = now_kst.strftime("%Y%m%d")
file_name = f"DRAM_논문_특허_GitHub_{today_ymd}.docx"

# =========================
# 유틸
# =========================
def parse_arxiv_date(date_str):
    return datetime.datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%SZ")


def build_google_scholar_url(query: str) -> str:
    return "https://scholar.google.com/scholar?q=" + urllib.parse.quote(query)


def build_company_fallback(company: str) -> str:
    return build_google_scholar_url(f"{company} DRAM HBM DDR memory")


def normalize_text(*parts):
    return " ".join([p or "" for p in parts]).lower()


def summarize_paper_en(title: str, summary: str) -> str:
    """
    영어 초록을 바탕으로 영문 요약 생성
    """
    text = re.sub(r"\s+", " ", (summary or "").strip())
    title_l = (title or "").lower()
    text_l = text.lower()

    if not text:
        return "No abstract information is available."

    topic = "This paper studies a memory-system-related problem."
    contribution = "It presents a technical approach or analysis."
    impact = "It aims to improve performance, efficiency, or memory behavior."

    if "hbm" in title_l or "high bandwidth memory" in text_l:
        topic = "This paper focuses on HBM-related architecture or optimization."
    elif "dram" in title_l or "dram" in text_l:
        topic = "This paper focuses on DRAM architecture, usage, or optimization."
    elif "ddr5" in title_l or "ddr5" in text_l:
        topic = "This paper discusses DDR5-related architecture or system issues."
    elif "ddr6" in title_l or "ddr6" in text_l:
        topic = "This paper discusses DDR6-related architecture or system issues."
    elif "lpddr5" in title_l or "lpddr5" in text_l:
        topic = "This paper discusses LPDDR5-related architecture or system issues."
    elif "lpddr6" in title_l or "lpddr6" in text_l:
        topic = "This paper discusses LPDDR6-related architecture or system issues."
    elif "pim" in title_l or "processing in memory" in text_l or "compute-in-memory" in text_l or "cim" in title_l:
        topic = "This paper studies a PIM/CIM-based memory acceleration approach."
    elif "llm" in title_l or "large language model" in text_l:
        topic = "This paper studies memory bottlenecks or acceleration for LLM inference."

    if any(k in text_l for k in ["propose", "proposes", "proposed", "introduce", "introduces", "present", "presents"]):
        contribution = "It proposes a new mechanism, architecture, or method."
    if any(k in text_l for k in ["architecture", "framework", "infrastructure", "prototype", "system"]):
        contribution = "It presents a system architecture, framework, or prototype design."
    if any(k in text_l for k in ["co-design", "codesign"]):
        contribution = "It adopts a hardware-software co-design approach."
    if any(k in text_l for k in ["evaluation", "evaluate", "analysis", "study"]):
        contribution = "It provides evaluation, comparison, or analysis results."

    if any(k in text_l for k in ["energy efficiency", "energy-efficient", "power efficiency", "efficient"]):
        impact = "It aims to improve power or energy efficiency."
    if any(k in text_l for k in ["latency", "throughput", "performance", "bandwidth"]):
        impact = "It aims to improve latency, throughput, performance, or bandwidth."
    if any(k in text_l for k in ["memory-intensive", "memory bottleneck", "memory wall"]):
        impact = "It aims to reduce memory bottlenecks."
    if any(k in text_l for k in ["edge", "edge device", "edge npu"]):
        impact = "It targets resource-constrained edge environments."

    return f"{topic} {contribution} {impact}"


def translate_summary_to_korean(english_summary: str) -> str:
    """
    영문 요약을 규칙 기반으로 한글 번역
    """
    if not english_summary:
        return "요약 정보 없음"

    translated = english_summary

    replacements = [
        ("This paper focuses on HBM-related architecture or optimization.", "이 논문은 HBM 관련 구조 또는 최적화를 다룬다."),
        ("This paper focuses on DRAM architecture, usage, or optimization.", "이 논문은 DRAM 구조, 활용, 또는 최적화를 다룬다."),
        ("This paper discusses DDR5-related architecture or system issues.", "이 논문은 DDR5 관련 구조 또는 시스템 이슈를 다룬다."),
        ("This paper discusses DDR6-related architecture or system issues.", "이 논문은 DDR6 관련 구조 또는 시스템 이슈를 다룬다."),
        ("This paper discusses LPDDR5-related architecture or system issues.", "이 논문은 LPDDR5 관련 구조 또는 시스템 이슈를 다룬다."),
        ("This paper discusses LPDDR6-related architecture or system issues.", "이 논문은 LPDDR6 관련 구조 또는 시스템 이슈를 다룬다."),
        ("This paper studies a PIM/CIM-based memory acceleration approach.", "이 논문은 PIM/CIM 기반 메모리 가속 구조를 다룬다."),
        ("This paper studies memory bottlenecks or acceleration for LLM inference.", "이 논문은 LLM 추론 과정의 메모리 병목 또는 가속 구조를 다룬다."),
        ("This paper studies a memory-system-related problem.", "이 논문은 메모리 시스템 관련 문제를 다룬다."),
        ("It proposes a new mechanism, architecture, or method.", "새로운 메커니즘, 구조, 또는 기법을 제안한다."),
        ("It presents a system architecture, framework, or prototype design.", "시스템 구조, 프레임워크, 또는 프로토타입 설계를 제시한다."),
        ("It adopts a hardware-software co-design approach.", "하드웨어-소프트웨어 공동 설계 접근을 적용한다."),
        ("It provides evaluation, comparison, or analysis results.", "평가, 비교, 또는 분석 결과를 제시한다."),
        ("It presents a technical approach or analysis.", "기술적 접근 또는 분석 내용을 제시한다."),
        ("It aims to improve power or energy efficiency.", "전력 또는 에너지 효율 향상을 목표로 한다."),
        ("It aims to improve latency, throughput, performance, or bandwidth.", "지연시간, 처리량, 성능, 또는 대역폭 개선을 목표로 한다."),
        ("It aims to reduce memory bottlenecks.", "메모리 병목 완화를 목표로 한다."),
        ("It targets resource-constrained edge environments.", "자원이 제한된 엣지 환경을 대상으로 한다."),
        ("It aims to improve performance, efficiency, or memory behavior.", "성능, 효율, 또는 메모리 동작 개선을 목표로 한다."),
        ("No abstract information is available.", "초록 정보가 없다."),
    ]

    for en, ko in replacements:
        translated = translated.replace(en, ko)

    return translated


def summarize_paper(title: str, summary: str) -> dict:
    summary_en = summarize_paper_en(title, summary)
    summary_ko = translate_summary_to_korean(summary_en)
    return {"en": summary_en, "ko": summary_ko}


def validate_arxiv_abs_url(url: str) -> bool:
    """
    실제 arXiv abs 링크 형식인지 검수
    """
    if not url:
        return False
    return bool(re.match(r"^https?://arxiv\.org/abs/[A-Za-z0-9.\-]+(v\d+)?$", url.strip()))


def get_best_link(paper):
    """
    arXiv와 Google Scholar만 사용
    arXiv 링크가 검증되면 우선 사용, 아니면 Scholar 검색 링크 사용
    """
    if paper.get("abs_url") and validate_arxiv_abs_url(paper["abs_url"]):
        return "arXiv", paper["abs_url"]
    return "Google Scholar", paper["scholar_url"]


def require_env(name, value):
    if not value:
        raise RuntimeError(f"{name} is missing or empty")


# =========================
# DOCX 하이퍼링크
# =========================
def add_hyperlink(paragraph, text, url, color="0563C1", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_paper_block(doc, idx, paper):
    doc.add_paragraph(f"{idx}) {paper['title']}")
    doc.add_paragraph(f"   - 저자: {paper['authors'] or '정보 없음'}")

    p = doc.add_paragraph()
    p.add_run(f"   - 대표 링크 ({paper['best_link_type']}): ")
    add_hyperlink(p, paper["best_link_url"], paper["best_link_url"])

    p = doc.add_paragraph()
    p.add_run("   - Google Scholar: ")
    add_hyperlink(p, paper["scholar_url"], paper["scholar_url"])

    doc.add_paragraph(f"   - 영문 요약: {paper['summary_en']}")
    doc.add_paragraph(f"   - 한글 요약: {paper['summary_ko']}")


def add_company_block(doc, company, item):
    doc.add_paragraph(company)
    doc.add_paragraph(f"   - 상태: {item['status']}")
    doc.add_paragraph(f"   - 제목/설명: {item['title']}")

    p = doc.add_paragraph()
    p.add_run(f"   - 대표 링크 ({item['link_type']}): ")
    add_hyperlink(p, item["url"], item["url"])

    doc.add_paragraph(f"   - 비고: {item['note']}")


def add_venue_block(doc, paper):
    doc.add_paragraph(f"- {paper['title']}")
    p = doc.add_paragraph()
    p.add_run(f"  · 링크 ({paper['best_link_type']}): ")
    add_hyperlink(p, paper["best_link_url"], paper["best_link_url"])


# =========================
# arXiv 검색
# =========================
def search_arxiv():
    query = " OR ".join([f'all:\"{term}\"' for term in SEARCH_TERMS])
    api_url = (
        "http://export.arxiv.org/api/query?"
        + urllib.parse.urlencode({
            "search_query": query,
            "start": 0,
            "max_results": ARXIV_MAX_RESULTS,
            "sortBy": "submittedDate",
            "sortOrder": "descending"
        })
    )

    response = requests.get(api_url, timeout=30)
    response.raise_for_status()
    feed = feedparser.parse(response.text)

    all_papers = []
    recent_1d_papers = []

    for entry in feed.entries:
        title = entry.title.strip().replace("\n", " ")
        summary = (entry.summary or "").strip().replace("\n", " ")
        published = entry.published
        published_dt = parse_arxiv_date(published)
        authors = ", ".join(a.name for a in entry.authors) if hasattr(entry, "authors") else ""
        abs_url = entry.id.strip()

        scholar_url = build_google_scholar_url(title)

        record = {
            "title": title,
            "summary": summary,
            "authors": authors,
            "published_dt": published_dt,
            "abs_url": abs_url if validate_arxiv_abs_url(abs_url) else "",
            "scholar_url": scholar_url,
        }

        record["best_link_type"], record["best_link_url"] = get_best_link(record)

        summary_obj = summarize_paper(title, summary)
        record["summary_en"] = summary_obj["en"]
        record["summary_ko"] = summary_obj["ko"]

        all_papers.append(record)

        if published_dt >= cutoff_utc:
            recent_1d_papers.append(record)

    return all_papers, recent_1d_papers


# =========================
# 회사별 항목
# =========================
def build_company_items(all_papers):
    company_hits = defaultdict(list)

    for paper in all_papers:
        hay = normalize_text(paper["title"], paper["summary"], paper["authors"])
        for company, kws in COMPANY_KEYWORDS.items():
            if any(kw in hay for kw in kws):
                company_hits[company].append(paper)

    company_items = {}
    for company in COMPANY_KEYWORDS.keys():
        if company_hits[company]:
            best = sorted(company_hits[company], key=lambda x: x["published_dt"], reverse=True)[0]
            company_items[company] = {
                "status": "확인",
                "title": best["title"],
                "link_type": best["best_link_type"],
                "url": best["best_link_url"],
                "note": "자동 검색 결과에서 회사 관련 항목 탐지",
            }
        else:
            company_items[company] = {
                "status": "fallback",
                "title": f"{company} 관련 DRAM/HBM 연구 검색",
                "link_type": "Google Scholar",
                "url": build_company_fallback(company),
                "note": "최근 1일/자동 검색에서 직접 논문 미확인, 대체 검색 링크 제공",
            }

    return company_items


# =========================
# DOCX 생성
# =========================
def create_docx(all_papers, recent_1d_papers, company_items, path):
    fallback_papers = all_papers[:5]

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Malgun Gothic"
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    doc.add_heading(f"Daily DRAM Report ({today_ymd})", 0)
    doc.add_paragraph("실행 시각: 07:07 (Asia/Seoul 목표)")
    doc.add_paragraph("검색 범위: 최근 1일 DRAM / HBM / DDR5 / DDR6 / LPDDR5 / LPDDR6 관련 논문")
    doc.add_paragraph("검색 소스: arXiv 자동 검색 + Google Scholar 검색 링크")

    doc.add_heading("0. 오늘의 요약", level=1)
    doc.add_paragraph(f"- 최근 1일 신규 논문 수: {len(recent_1d_papers)}건")
    doc.add_paragraph(f"- 자동 검색 전체 논문 수: {len(all_papers)}건")
    doc.add_paragraph("- 필수 기업 4개는 별도 섹션에서 확인 결과 또는 대체 검색 링크를 제공")

    doc.add_heading("1. 최근 1일 논문 (엄격 기준)", level=1)
    if recent_1d_papers:
        for idx, paper in enumerate(recent_1d_papers[:5], 1):
            add_paper_block(doc, idx, paper)
    else:
        doc.add_paragraph("신규 공개 논문 없음")

    doc.add_heading("2. 참고: 최근 공개 논문 (1일 초과)", level=1)
    if fallback_papers:
        for idx, paper in enumerate(fallback_papers, 1):
            add_paper_block(doc, idx, paper)
    else:
        doc.add_paragraph("참고 논문 없음")

    doc.add_heading("3. 필수 기업 포함 항목", level=1)
    for company, item in company_items.items():
        add_company_block(doc, company, item)

    doc.add_heading("4. 참고 링크", level=1)
    shown = 0
    for paper in all_papers:
        add_venue_block(doc, paper)
        shown += 1
        if shown >= 5:
            break
    if shown == 0:
        doc.add_paragraph("자동 탐지 결과 없음")

    doc.add_heading("5. 메모", level=1)
    doc.add_paragraph("- 링크는 arXiv와 Google Scholar만 유지했다.")
    doc.add_paragraph("- DOI, 공식 페이지, 공개시각은 문서에서 제거했다.")
    doc.add_paragraph("- arXiv 링크는 실제 API 응답의 abs URL만 사용해 검수했다.")
    doc.add_paragraph("- Google Scholar는 검색 링크만 제공하며, 가상 논문 링크는 생성하지 않았다.")
    doc.add_paragraph("- 문서 내 URL은 클릭 가능한 하이퍼링크로 삽입했다.")
    doc.add_paragraph("- 요약은 영문 요약 생성 후 한글 번역을 추가하는 방식으로 구성했다.")

    doc.save(path)


# =========================
# Drive 업로드
# =========================
def get_drive_service():
    require_env("GOOGLE_CLIENT_ID", GOOGLE_CLIENT_ID)
    require_env("GOOGLE_CLIENT_SECRET", GOOGLE_CLIENT_SECRET)
    require_env("GOOGLE_REFRESH_TOKEN", GOOGLE_REFRESH_TOKEN)
    require_env("GOOGLE_DRIVE_FOLDER_ID", GOOGLE_DRIVE_FOLDER_ID)

    creds = Credentials(
        None,
        refresh_token=GOOGLE_REFRESH_TOKEN,
        token_uri="https://oauth2.googleapis.com/token",
        client_id=GOOGLE_CLIENT_ID,
        client_secret=GOOGLE_CLIENT_SECRET,
        scopes=["https://www.googleapis.com/auth/drive.file"],
    )
    creds.refresh(GoogleRequest())
    return build("drive", "v3", credentials=creds)


def upload_to_drive(path):
    service = get_drive_service()

    media = MediaFileUpload(
        path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    file_metadata = {
        "name": os.path.basename(path),
        "parents": [GOOGLE_DRIVE_FOLDER_ID]
    }

    created = service.files().create(
        body=file_metadata,
        media_body=media,
        fields="id,name,webViewLink"
    ).execute()

    print("Uploaded:", created["name"])
    print("Link:", created.get("webViewLink"))


# =========================
# main
# =========================
def main():
    all_papers, recent_1d_papers = search_arxiv()
    company_items = build_company_items(all_papers)
    create_docx(all_papers, recent_1d_papers, company_items, file_name)
    upload_to_drive(file_name)


if __name__ == "__main__":
    main()
